#!/usr/bin/env python3
"""
Create final internship report following CB20114_FINAL REPORT.pdf format exactly
- Add UMP logo
- Match exact formatting
- Fill in blank spaces
- Proper page layout
"""

import re
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE

def add_page_border(section):
    """Add thin black border around the page content"""
    try:
        pgBorders = OxmlElement('w:pgBorders')
        pgBorders.set(qn('w:rsidR'), '00C47AE6')
        pgBorders.set(qn('w:rsidRPr'), '00C47AE6')
        pgBorders.set(qn('w:rsidRDefault'), '00C47AE6')
        pgBorders.set(qn('w:rsidP'), '00C47AE6')
        
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')  # Thin border
            border.set(qn('w:space'), '24')
            border.set(qn('w:color'), '000000')  # Black
            pgBorders.append(border)
        
        section._sectPr.append(pgBorders)
    except Exception as e:
        print(f"Warning: Could not add page border: {e}")

def add_page_number(section):
    """Add page number at bottom center"""
    try:
        footer = section.footer
        
        # Clear existing footer content
        if len(footer.paragraphs) > 0:
            for para in footer.paragraphs:
                footer._element.remove(para._element)
        
        paragraph = footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add page number field
        run = paragraph.add_run()
        
        # Create field code for page number
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._element.append(fldChar1)
        run._element.append(instrText)
        run._element.append(fldChar2)
        
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    except Exception as e:
        print(f"Warning: Could not add page number: {e}")

def add_ump_logo(doc):
    """Add UMP logo to cover page if available"""
    logo_paths = [
        'logomnsa1.jpeg',
        'public/logomnsa1.jpeg',
        'UMP_logo.png',
        'UMP_logo.jpg',
        'logo.png',
        'logo.jpg'
    ]
    
    for logo_path in logo_paths:
        if os.path.exists(logo_path):
            try:
                # Add logo paragraph at the very beginning
                if len(doc.paragraphs) > 0:
                    logo_para = doc.paragraphs[0].insert_paragraph_before()
                else:
                    logo_para = doc.add_paragraph()
                logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = logo_para.add_run()
                run.add_picture(logo_path, width=Inches(2.5))
                # Add spacing after logo
                spacing_para = doc.add_paragraph()
                spacing_para.paragraph_format.space_after = Pt(12)
                return True
            except Exception as e:
                print(f"Warning: Could not add logo from {logo_path}: {e}")
                continue
    
    print("Note: UMP logo not found. Please add manually to cover page.")
    return False

def clean_markdown_text(text):
    """Remove markdown formatting from text"""
    if not text:
        return ""
    # Remove bold
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    # Remove italic
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    # Remove code
    text = re.sub(r'`(.*?)`', r'\1', text)
    # Remove links but keep text
    text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
    # Convert line breaks
    text = text.replace('<br>', '\n')
    # Remove HTML tags
    text = re.sub(r'<[^>]+>', '', text)
    # Replace placeholders - keep placeholders for manual filling
    text = text.replace('[DATE]', '[To be filled]')
    text = text.replace('[SUPERVISOR NAME]', '[SUPERVISOR NAME - To be filled]')
    text = text.replace('[To be filled]', '[To be filled]')  # Keep as is
    return text.strip()

def process_table(lines, start_idx):
    """Process a markdown table and return table data and next index"""
    table_lines = []
    idx = start_idx
    
    while idx < len(lines) and lines[idx].strip().startswith('|'):
        table_lines.append(lines[idx].strip())
        idx += 1
    
    if len(table_lines) < 2:
        return None, idx
    
    # Parse table
    headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
    if not headers or headers[0] == '':
        return None, idx
    
    # Skip separator line
    data_rows = []
    for line in table_lines[2:]:
        if '---' not in line and line.strip():
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if len(cells) == len(headers):
                data_rows.append(cells)
    
    return {'headers': headers, 'rows': data_rows}, idx

def setup_document_styles(doc):
    """Setup document styles to match template format"""
    # Set default font to Times New Roman 12pt
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Set single line spacing for Normal style
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    
    # Setup heading styles
    for level in range(1, 5):
        try:
            heading_style = doc.styles[f'Heading {level}']
            heading_style.font.name = 'Times New Roman'
            heading_style.font.bold = True
            heading_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            heading_style.paragraph_format.line_spacing = 1.0
            
            if level == 1:
                heading_style.font.size = Pt(14)
                heading_style.paragraph_format.space_before = Pt(12)
                heading_style.paragraph_format.space_after = Pt(6)
            elif level == 2:
                heading_style.font.size = Pt(13)
                heading_style.paragraph_format.space_before = Pt(12)
                heading_style.paragraph_format.space_after = Pt(6)
            else:
                heading_style.font.size = Pt(12)
                heading_style.paragraph_format.space_before = Pt(12)
                heading_style.paragraph_format.space_after = Pt(6)
        except:
            pass

def create_word_document():
    """Create Word document from markdown file with exact template format"""
    # Read markdown file
    try:
        with open('INDUSTRIAL_TRAINING_REPORT.md', 'r', encoding='utf-8') as f:
            content = f.read()
    except FileNotFoundError:
        print("Error: INDUSTRIAL_TRAINING_REPORT.md not found.")
        return
    
    # Create Word document
    doc = Document()
    
    # Set page margins (2.5cm left/right, 3.0cm top, 2.5cm bottom) - matching PDF format
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(3.0)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        
        # Add page border
        add_page_border(section)
    
    # Setup document styles
    setup_document_styles(doc)
    
    # Add UMP logo to cover page
    add_ump_logo(doc)
    
    # Process content line by line
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Handle headings
        if line.startswith('# '):
            # Main heading (H1)
            heading_text = line[2:].strip()
            p = doc.add_heading(heading_text, level=1)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
                run.bold = True
        elif line.startswith('## '):
            # Subheading (H2)
            heading_text = line[3:].strip()
            p = doc.add_heading(heading_text, level=2)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(13)
                run.bold = True
        elif line.startswith('### '):
            # Sub-subheading (H3)
            heading_text = line[4:].strip()
            p = doc.add_heading(heading_text, level=3)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.bold = True
        elif line.startswith('#### '):
            # Level 4 heading
            heading_text = line[5:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(heading_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = True
        elif line.startswith('|') and '|' in line:
            # Table - process it
            table_data, next_idx = process_table(lines, i)
            if table_data:
                # Create table
                num_cols = len(table_data['headers'])
                num_rows = len(table_data['rows']) + 1  # +1 for header
                
                table = doc.add_table(rows=num_rows, cols=num_cols)
                table.style = 'Light Grid Accent 1'
                
                # Add header row
                header_cells = table.rows[0].cells
                for j, header in enumerate(table_data['headers']):
                    header_cells[j].text = clean_markdown_text(header)
                    for paragraph in header_cells[j].paragraphs:
                        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                        paragraph.paragraph_format.line_spacing = 1.0
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(12)
                            run.bold = True
                
                # Add data rows
                for row_idx, row_data in enumerate(table_data['rows']):
                    row = table.rows[row_idx + 1]
                    for j, cell_data in enumerate(row_data):
                        if j < len(row.cells):
                            row.cells[j].text = clean_markdown_text(cell_data)
                            for paragraph in row.cells[j].paragraphs:
                                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                                paragraph.paragraph_format.line_spacing = 1.0
                                for run in paragraph.runs:
                                    run.font.name = 'Times New Roman'
                                    run.font.size = Pt(12)
                
                i = next_idx - 1  # Adjust for loop increment
        elif line.startswith('---'):
            # Horizontal rule - add minimal spacing
            doc.add_paragraph()
        else:
            # Regular paragraph
            # Check if it's a list item
            if line.startswith('- ') or line.startswith('* '):
                # Bullet list
                list_text = line[2:].strip()
                p = doc.add_paragraph(clean_markdown_text(list_text), style='List Bullet')
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
            elif re.match(r'^\d+\.\s', line):
                # Numbered list
                list_text = re.sub(r'^\d+\.\s', '', line)
                p = doc.add_paragraph(clean_markdown_text(list_text), style='List Number')
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
            else:
                # Regular paragraph with single spacing
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(clean_markdown_text(line))
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
        
        i += 1
    
    # Add page numbers to all sections
    for section in doc.sections:
        add_page_number(section)
    
    # Save document
    output_file = 'INDUSTRIAL_TRAINING_REPORT_FINAL.docx'
    doc.save(output_file)
    print(f"✓ Final Word document created successfully: {output_file}")
    print(f"✓ File saved in: {os.path.abspath(output_file)}")
    print("\nFormatting applied:")
    print("  ✓ Thin black border around content")
    print("  ✓ Single line spacing (1.0)")
    print("  ✓ Page numbers at bottom center")
    print("  ✓ Times New Roman 12pt font")
    print("  ✓ Proper margins (2.5cm left/right, 3.0cm top, 2.5cm bottom)")
    print("  ✓ Proper heading styles")
    print("\nNote: Please review the document and:")
    print("  1. Verify UMP logo is added to cover page")
    print("  2. Fill in supervisor name and date in confidentiality section")
    print("  3. Add organization chart in attachment section")
    print("  4. Format Gantt chart table with colors/shading if needed")

if __name__ == '__main__':
    try:
        create_word_document()
    except Exception as e:
        print(f"Error creating Word document: {e}")
        import traceback
        traceback.print_exc()

