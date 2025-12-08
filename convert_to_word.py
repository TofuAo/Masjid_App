#!/usr/bin/env python3
"""
Script to convert INDUSTRIAL_TRAINING_REPORT.md to Word document with proper formatting
"""

import re
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_word_document():
    # Read markdown file
    with open('INDUSTRIAL_TRAINING_REPORT.md', 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create Word document
    doc = Document()
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(3.0)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # Set default font to Times New Roman 12pt
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Process content
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
        
        # Handle headings
        if line.startswith('# '):
            # Main heading
            p = doc.add_heading(line[2:], level=1)
            p.runs[0].font.name = 'Times New Roman'
            p.runs[0].font.size = Pt(14)
            p.runs[0].bold = True
        elif line.startswith('## '):
            # Subheading
            p = doc.add_heading(line[3:], level=2)
            p.runs[0].font.name = 'Times New Roman'
            p.runs[0].font.size = Pt(13)
            p.runs[0].bold = True
        elif line.startswith('### '):
            # Sub-subheading
            p = doc.add_heading(line[4:], level=3)
            p.runs[0].font.name = 'Times New Roman'
            p.runs[0].font.size = Pt(12)
            p.runs[0].bold = True
        elif line.startswith('#### '):
            # Level 4 heading
            p = doc.add_paragraph(line[5:])
            p.runs[0].font.name = 'Times New Roman'
            p.runs[0].font.size = Pt(12)
            p.runs[0].bold = True
        elif line.startswith('|') and '|' in line:
            # Table row - handle tables
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            i -= 1  # Adjust for loop increment
            
            if len(table_lines) > 1:
                # Create table
                headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
                if len(headers) > 0:
                    table = doc.add_table(rows=1, cols=len(headers))
                    table.style = 'Light Grid Accent 1'
                    
                    # Add header row
                    header_cells = table.rows[0].cells
                    for j, header in enumerate(headers):
                        header_cells[j].text = header
                        for paragraph in header_cells[j].paragraphs:
                            paragraph.runs[0].font.name = 'Times New Roman'
                            paragraph.runs[0].font.size = Pt(12)
                            paragraph.runs[0].bold = True
                    
                    # Add data rows
                    for row_line in table_lines[2:]:  # Skip header and separator
                        if '---' not in row_line:
                            cells = [cell.strip() for cell in row_line.split('|')[1:-1]]
                            if len(cells) == len(headers):
                                row = table.add_row()
                                for j, cell in enumerate(cells):
                                    # Remove markdown formatting
                                    cell = re.sub(r'\*\*(.*?)\*\*', r'\1', cell)  # Bold
                                    cell = re.sub(r'<br>', '\n', cell)  # Line breaks
                                    row.cells[j].text = cell
                                    for paragraph in row.cells[j].paragraphs:
                                        for run in paragraph.runs:
                                            run.font.name = 'Times New Roman'
                                            run.font.size = Pt(12)
        elif line.startswith('---'):
            # Horizontal rule - add page break or spacing
            doc.add_paragraph()
        else:
            # Regular paragraph
            p = doc.add_paragraph()
            # Remove markdown formatting
            text = line
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Bold
            text = re.sub(r'\*(.*?)\*', r'\1', text)  # Italic
            text = re.sub(r'`(.*?)`', r'\1', text)  # Code
            text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)  # Links
            text = re.sub(r'<br>', '\n', text)  # Line breaks
            
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        
        i += 1
    
    # Save document
    doc.save('INDUSTRIAL_TRAINING_REPORT.docx')
    print("Word document created successfully: INDUSTRIAL_TRAINING_REPORT.docx")

if __name__ == '__main__':
    try:
        create_word_document()
    except ImportError:
        print("Error: python-docx library not installed.")
        print("Please install it using: pip install python-docx")
    except Exception as e:
        print(f"Error creating Word document: {e}")
        print("\nAlternative: Use online markdown to Word converter or copy-paste to Word manually.")

