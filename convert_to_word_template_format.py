#!/usr/bin/env python3
"""
Enhanced script to convert INDUSTRIAL_TRAINING_REPORT.md to Word document
Following the format template: FORMAT FOR FINAL REPORT (SECOND EVALUATION)
- Times New Roman 12pt font
- Single line spacing (1.0)
- Thin black page borders
- Page numbers at bottom center
- Proper margins (2.5cm left/right, 3.0cm top, 2.5cm bottom)
- Proper heading styles
- Table formatting
"""

import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE

def add_page_border(section):
    """Add thin black border around the page content"""
    try:
        pgBorders = OxmlElement('w:pgBorders')
        pgBorders.set(qn('w:rsidR'), '00C47AE6')
        pgBorders.set(qn('w:rsidRPr'), '00C47AE6')
        pgBorders.set(qn('w:rsidRDefault'), '00C47AE6')
        pgBorders.set(qn('w:rsidP'), '00C47AE6')
        
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')  # Thin border (0.5pt)
            border.set(qn('w:space'), '24')
            border.set(qn('w:color'), '000000')  # Black
            pgBorders.append(border)
        
        section._sectPr.append(pgBorders)
    except Exception as e:
        print(f"Warning: Could not add page border: {e}")

def add_page_number(section):
    """Add page number at bottom center"""
    try:
        footer = section.footer
        
        # Clear existing footer content
        if len(footer.paragraphs) > 0:
            for para in footer.paragraphs:
                footer._element.remove(para._element)
        
        paragraph = footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add page number field
        run = paragraph.add_run()
        
        # Create field code for page number
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._element.append(fldChar1)
        run._element.append(instrText)
        run._element.append(fldChar2)
        
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    except Exception as e:
        print(f"Warning: Could not add page number: {e}")

def clean_markdown_text(text):
    """Remove markdown formatting from text"""
    if not text:
        return ""
    # Remove bold
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    # Remove italic
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    # Remove code
    text = re.sub(r'`(.*?)`', r'\1', text)
    # Remove links but keep text
    text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
    # Convert line breaks
    text = text.replace('<br>', '\n')
    # Remove HTML tags
    text = re.sub(r'<[^>]+>', '', text)
    return text.strip()

def setup_document_styles(doc):
    """Setup document styles to match template format"""
    # Set default font to Times New Roman 12pt
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Set single line spacing for Normal style
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    
    # Setup heading styles
    for level in range(1, 5):
        try:
            heading_style = doc.styles[f'Heading {level}']
            heading_style.font.name = 'Times New Roman'
            heading_style.font.bold = True
            heading_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            heading_style.paragraph_format.line_spacing = 1.0
            
            if level == 1:
                heading_style.font.size = Pt(14)
                heading_style.paragraph_format.space_before = Pt(12)
                heading_style.paragraph_format.space_after = Pt(6)
            elif level == 2:
                heading_style.font.size = Pt(13)
                heading_style.paragraph_format.space_before = Pt(12)
                heading_style.paragraph_format.space_after = Pt(6)
            else:
                heading_style.font.size = Pt(12)
                heading_style.paragraph_format.space_before = Pt(12)
                heading_style.paragraph_format.space_after = Pt(6)
        except:
            pass

def process_table(lines, start_idx):
    """Process a markdown table and return table data and next index"""
    table_lines = []
    idx = start_idx
    
    while idx < len(lines) and lines[idx].strip().startswith('|'):
        table_lines.append(lines[idx].strip())
        idx += 1
    
    if len(table_lines) < 2:
        return None, idx
    
    # Parse table
    headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
    if not headers or headers[0] == '':
        return None, idx
    
    # Skip separator line
    data_rows = []
    for line in table_lines[2:]:
        if '---' not in line and line.strip():
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if len(cells) == len(headers):
                data_rows.append(cells)
    
    return {'headers': headers, 'rows': data_rows}, idx

def create_word_document():
    """Create Word document from markdown file with exact template format"""
    # Read markdown file
    try:
        with open('INDUSTRIAL_TRAINING_REPORT.md', 'r', encoding='utf-8') as f:
            content = f.read()
    except FileNotFoundError:
        print("Error: INDUSTRIAL_TRAINING_REPORT.md not found.")
        print("Please make sure the file exists in the current directory.")
        return
    
    # Create Word document
    doc = Document()
    
    # Set page margins (2.5cm left/right, 3.0cm top, 2.5cm bottom)
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(3.0)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        
        # Add page border
        add_page_border(section)
    
    # Setup document styles
    setup_document_styles(doc)
    
    # Process content line by line
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Handle headings
        if line.startswith('# '):
            # Main heading (H1)
            heading_text = line[2:].strip()
            p = doc.add_heading(heading_text, level=1)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
                run.bold = True
        elif line.startswith('## '):
            # Subheading (H2)
            heading_text = line[3:].strip()
            p = doc.add_heading(heading_text, level=2)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(13)
                run.bold = True
        elif line.startswith('### '):
            # Sub-subheading (H3)
            heading_text = line[4:].strip()
            p = doc.add_heading(heading_text, level=3)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.bold = True
        elif line.startswith('#### '):
            # Level 4 heading
            heading_text = line[5:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(heading_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = True
        elif line.startswith('|') and '|' in line:
            # Table - process it
            table_data, next_idx = process_table(lines, i)
            if table_data:
                # Create table
                num_cols = len(table_data['headers'])
                num_rows = len(table_data['rows']) + 1  # +1 for header
                
                table = doc.add_table(rows=num_rows, cols=num_cols)
                table.style = 'Light Grid Accent 1'
                
                # Add header row
                header_cells = table.rows[0].cells
                for j, header in enumerate(table_data['headers']):
                    header_cells[j].text = clean_markdown_text(header)
                    for paragraph in header_cells[j].paragraphs:
                        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                        paragraph.paragraph_format.line_spacing = 1.0
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(12)
                            run.bold = True
                
                # Add data rows
                for row_idx, row_data in enumerate(table_data['rows']):
                    row = table.rows[row_idx + 1]
                    for j, cell_data in enumerate(row_data):
                        if j < len(row.cells):
                            row.cells[j].text = clean_markdown_text(cell_data)
                            for paragraph in row.cells[j].paragraphs:
                                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                                paragraph.paragraph_format.line_spacing = 1.0
                                for run in paragraph.runs:
                                    run.font.name = 'Times New Roman'
                                    run.font.size = Pt(12)
                
                i = next_idx - 1  # Adjust for loop increment
        elif line.startswith('---'):
            # Horizontal rule - add minimal spacing
            doc.add_paragraph()
        else:
            # Regular paragraph
            # Check if it's a list item
            if line.startswith('- ') or line.startswith('* '):
                # Bullet list
                list_text = line[2:].strip()
                p = doc.add_paragraph(clean_markdown_text(list_text), style='List Bullet')
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
            elif re.match(r'^\d+\.\s', line):
                # Numbered list
                list_text = re.sub(r'^\d+\.\s', '', line)
                p = doc.add_paragraph(clean_markdown_text(list_text), style='List Number')
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
            else:
                # Regular paragraph with single spacing
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(clean_markdown_text(line))
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
        
        i += 1
    
    # Add page numbers to all sections
    for section in doc.sections:
        add_page_number(section)
    
    # Save document
    output_file = 'INDUSTRIAL_TRAINING_REPORT_UPDATED.docx'
    doc.save(output_file)
    print(f"✓ Word document created successfully: {output_file}")
    print(f"✓ File saved in: {os.path.abspath(output_file)}")
    print("\nFormatting applied:")
    print("  ✓ Thin black border around content")
    print("  ✓ Single line spacing (1.0)")
    print("  ✓ Page numbers at bottom center")
    print("  ✓ Times New Roman 12pt font")
    print("  ✓ Proper margins (2.5cm left/right, 3.0cm top, 2.5cm bottom)")
    print("  ✓ Proper heading styles")
    print("\nNote: Please review the document and:")
    print("  1. Verify all formatting matches the template")
    print("  2. Add UMP logo to cover page if required")
    print("  3. Format Gantt chart table with colors/shading if needed")
    print("  4. Add organization chart (if available)")
    print("  5. Add supervisor name and date in appropriate sections")

if __name__ == '__main__':
    try:
        import os
        create_word_document()
    except Exception as e:
        print(f"Error creating Word document: {e}")
        import traceback
        traceback.print_exc()

