#!/usr/bin/env python3
"""
Improved script to convert INDUSTRIAL_TRAINING_REPORT.md to Word document with proper formatting
"""

import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def clean_markdown_text(text):
    """Remove markdown formatting from text"""
    # Remove bold
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    # Remove italic
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    # Remove code
    text = re.sub(r'`(.*?)`', r'\1', text)
    # Remove links but keep text
    text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
    # Convert line breaks
    text = text.replace('<br>', '\n')
    # Remove HTML tags
    text = re.sub(r'<[^>]+>', '', text)
    return text.strip()

def add_formatted_paragraph(doc, text, font_name='Times New Roman', font_size=12, bold=False, alignment=None):
    """Add a formatted paragraph to the document"""
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    
    # Split by line breaks
    parts = text.split('\n')
    for i, part in enumerate(parts):
        if i > 0:
            p.add_run().add_break()
        run = p.add_run(clean_markdown_text(part))
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.bold = bold
    
    return p

def process_table(lines, start_idx):
    """Process a markdown table and return table data and next index"""
    table_lines = []
    idx = start_idx
    
    while idx < len(lines) and lines[idx].strip().startswith('|'):
        table_lines.append(lines[idx].strip())
        idx += 1
    
    if len(table_lines) < 2:
        return None, idx
    
    # Parse table
    headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
    if not headers or headers[0] == '':
        return None, idx
    
    # Skip separator line
    data_rows = []
    for line in table_lines[2:]:
        if '---' not in line and line.strip():
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if len(cells) == len(headers):
                data_rows.append(cells)
    
    return {'headers': headers, 'rows': data_rows}, idx

def create_word_document():
    """Create Word document from markdown file"""
    # Read markdown file
    with open('INDUSTRIAL_TRAINING_REPORT.md', 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create Word document
    doc = Document()
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(3.0)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # Set default font to Times New Roman 12pt
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Process content line by line
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines (but add spacing)
        if not line:
            i += 1
            continue
        
        # Handle headings
        if line.startswith('# '):
            # Main heading (H1)
            heading_text = line[2:].strip()
            p = doc.add_heading(heading_text, level=1)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
                run.bold = True
        elif line.startswith('## '):
            # Subheading (H2)
            heading_text = line[3:].strip()
            p = doc.add_heading(heading_text, level=2)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(13)
                run.bold = True
        elif line.startswith('### '):
            # Sub-subheading (H3)
            heading_text = line[4:].strip()
            p = doc.add_heading(heading_text, level=3)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.bold = True
        elif line.startswith('#### '):
            # Level 4 heading
            heading_text = line[5:].strip()
            p = doc.add_paragraph()
            run = p.add_run(heading_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = True
        elif line.startswith('|') and '|' in line:
            # Table - process it
            table_data, next_idx = process_table(lines, i)
            if table_data:
                # Create table
                num_cols = len(table_data['headers'])
                num_rows = len(table_data['rows']) + 1  # +1 for header
                
                table = doc.add_table(rows=num_rows, cols=num_cols)
                table.style = 'Light Grid Accent 1'
                
                # Add header row
                header_cells = table.rows[0].cells
                for j, header in enumerate(table_data['headers']):
                    header_cells[j].text = clean_markdown_text(header)
                    for paragraph in header_cells[j].paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(12)
                            run.bold = True
                
                # Add data rows
                for row_idx, row_data in enumerate(table_data['rows']):
                    row = table.rows[row_idx + 1]
                    for j, cell_data in enumerate(row_data):
                        if j < len(row.cells):
                            row.cells[j].text = clean_markdown_text(cell_data)
                            for paragraph in row.cells[j].paragraphs:
                                for run in paragraph.runs:
                                    run.font.name = 'Times New Roman'
                                    run.font.size = Pt(12)
                
                i = next_idx - 1  # Adjust for loop increment
        elif line.startswith('---'):
            # Horizontal rule - add spacing
            doc.add_paragraph()
        else:
            # Regular paragraph
            # Check if it's a list item
            if line.startswith('- ') or line.startswith('* '):
                # Bullet list
                list_text = line[2:].strip()
                p = doc.add_paragraph(clean_markdown_text(list_text), style='List Bullet')
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
            elif re.match(r'^\d+\.\s', line):
                # Numbered list
                list_text = re.sub(r'^\d+\.\s', '', line)
                p = doc.add_paragraph(clean_markdown_text(list_text), style='List Number')
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
            else:
                # Regular paragraph
                p = add_formatted_paragraph(doc, line, font_name='Times New Roman', font_size=12)
        
        i += 1
    
    # Save document
    output_file = 'INDUSTRIAL_TRAINING_REPORT.docx'
    doc.save(output_file)
    print(f"✓ Word document created successfully: {output_file}")
    print(f"✓ File saved in: {__import__('os').path.abspath(output_file)}")
    print("\nNote: Please review the document and:")
    print("  1. Add UMP logo to cover page")
    print("  2. Format Gantt chart table with colors/shading")
    print("  3. Add organization chart (if available)")
    print("  4. Verify all formatting is correct")
    print("  5. Add page numbers")

if __name__ == '__main__':
    try:
        create_word_document()
    except FileNotFoundError:
        print("Error: INDUSTRIAL_TRAINING_REPORT.md not found.")
        print("Please make sure the file exists in the current directory.")
    except Exception as e:
        print(f"Error creating Word document: {e}")
        import traceback
        traceback.print_exc()

